#!/usr/bin/env python3
"""
Script para generar documentación Word del proyecto Arynstal.
Ejecutar desde el directorio raíz del proyecto con el venv activado.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def crear_estilos(doc):
    """Crea estilos personalizados para el documento."""
    # Los estilos por defecto de Word son suficientes para este documento
    pass

def agregar_portada(doc):
    """Agrega la portada del documento."""
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ARYNSTAL")
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(0, 102, 204)

    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Sistema CRM para Gestión de Instalaciones y Reformas")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tipo de documento
    tipo = doc.add_paragraph()
    tipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tipo.add_run("DOCUMENTACIÓN DEL PROCESO DE DESARROLLO")
    run.bold = True
    run.font.size = Pt(16)

    # Información adicional
    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Autor: @cgvrzon\n").bold = True
    info.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}\n")
    info.add_run("Versión: 1.0\n")
    info.add_run("Primera web - Proyecto de aprendizaje")

    doc.add_page_break()

def agregar_indice(doc):
    """Agrega el índice del documento."""
    doc.add_heading("ÍNDICE", level=1)

    indice = [
        ("1. Introducción al Proyecto", 3),
        ("   1.1. Descripción general", 3),
        ("   1.2. Objetivos", 3),
        ("   1.3. Stack tecnológico", 4),
        ("2. Análisis y Decisiones Iniciales", 5),
        ("   2.1. Análisis de requisitos", 5),
        ("   2.2. Arquitectura elegida", 5),
        ("   2.3. Decisiones de diseño", 6),
        ("3. Desarrollo del Código - Fases", 7),
        ("   3.1. Fase 1-3: Estructura base", 7),
        ("   3.2. Fase 4-6: Frontend y CRM", 8),
        ("   3.3. Fase 7-9: Formulario y seguridad", 9),
        ("4. Filosofías y Convenciones Aplicadas", 10),
        ("   4.1. Clean Code y SOLID", 10),
        ("   4.2. Otras filosofías web", 11),
        ("5. Documentación y Organización", 12),
        ("   5.1. Estructura de documentación", 12),
        ("   5.2. Comentarios en código", 12),
        ("6. Plan de Producción - Ruta Segura", 13),
        ("   6.1. Tareas críticas", 13),
        ("   6.2. Cronograma", 14),
        ("7. Lecciones Aprendidas", 15),
        ("8. Próximos Pasos", 16),
    ]

    for item, _ in indice:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

def agregar_introduccion(doc):
    """Agrega la sección de introducción."""
    doc.add_heading("1. Introducción al Proyecto", level=1)

    # 1.1 Descripción
    doc.add_heading("1.1. Descripción General", level=2)
    doc.add_paragraph(
        "Arynstal es un sistema CRM (Customer Relationship Management) desarrollado "
        "para una empresa de instalaciones y reformas en Barcelona. El proyecto permite "
        "gestionar leads (solicitudes de clientes potenciales), presupuestos, y servicios "
        "ofrecidos por la empresa."
    )
    doc.add_paragraph(
        "Este proyecto representa mi primera web completa, desarrollada con el objetivo "
        "de aprender las mejores prácticas de desarrollo web profesional, desde la "
        "arquitectura hasta el despliegue en producción."
    )

    # 1.2 Objetivos
    doc.add_heading("1.2. Objetivos del Proyecto", level=2)

    objetivos = [
        "Crear una web corporativa profesional para Arynstal SL",
        "Implementar un sistema de gestión de leads desde formulario de contacto",
        "Desarrollar un panel de administración completo para el equipo",
        "Aplicar buenas prácticas: Clean Code, SOLID, seguridad web",
        "Aprender el ciclo completo: desarrollo → testing → deployment",
        "Documentar todo el proceso para futuras referencias",
    ]

    for obj in objetivos:
        doc.add_paragraph(obj, style='List Bullet')

    # 1.3 Stack
    doc.add_heading("1.3. Stack Tecnológico", level=2)

    doc.add_paragraph("Backend:", style='List Bullet')
    p = doc.add_paragraph("Django 5.x/6.x (framework web Python)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("PostgreSQL (producción) / SQLite (desarrollo)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("Gunicorn (servidor WSGI)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph("Frontend:", style='List Bullet')
    p = doc.add_paragraph("HTML5 + Tailwind CSS", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("Vite (bundler de assets)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("JavaScript vanilla", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph("Infraestructura:", style='List Bullet')
    p = doc.add_paragraph("Hetzner Cloud CX22 (VPS)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("Nginx (reverse proxy)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("Let's Encrypt (SSL)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph("Cloudflare (DNS)", style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

def agregar_analisis(doc):
    """Agrega la sección de análisis y decisiones."""
    doc.add_heading("2. Análisis y Decisiones Iniciales", level=1)

    # 2.1 Requisitos
    doc.add_heading("2.1. Análisis de Requisitos", level=2)

    doc.add_paragraph("Requisitos funcionales identificados:")
    requisitos_func = [
        "Web corporativa con información de servicios",
        "Formulario de contacto que cree leads en el sistema",
        "Panel de administración para gestionar leads",
        "Sistema de presupuestos vinculados a leads",
        "Catálogo de servicios editable",
        "Sistema de usuarios con roles (admin, oficina, técnico)",
        "Historial de cambios en leads (auditoría)",
        "Notificaciones por email al recibir solicitudes",
    ]
    for req in requisitos_func:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("Requisitos no funcionales:")
    requisitos_no_func = [
        "Seguridad: Protección contra spam, CSRF, XSS",
        "Rendimiento: Carga < 3 segundos",
        "Mantenibilidad: Código documentado y testeable",
        "Escalabilidad: Preparado para crecer",
        "Cumplimiento RGPD: Privacidad y consentimiento",
    ]
    for req in requisitos_no_func:
        doc.add_paragraph(req, style='List Bullet')

    # 2.2 Arquitectura
    doc.add_heading("2.2. Arquitectura Elegida", level=2)

    doc.add_paragraph(
        "Se optó por una arquitectura monolítica con Django por ser un primer proyecto. "
        "Esta decisión se basa en:"
    )

    razones = [
        "Django incluye todo lo necesario (ORM, admin, auth, forms)",
        "Menor complejidad que microservicios",
        "Documentación abundante y comunidad activa",
        "Suficiente para el volumen esperado de tráfico",
        "Fácil de mantener por una sola persona",
    ]
    for razon in razones:
        doc.add_paragraph(razon, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("Estructura de apps Django:")

    apps = [
        "apps/leads/ - CRM central (Lead, Budget, LeadLog, LeadImage)",
        "apps/services/ - Catálogo de servicios",
        "apps/users/ - Perfiles de usuario con roles",
        "apps/web/ - Vistas públicas del frontend",
    ]
    for app in apps:
        p = doc.add_paragraph(app)
        p.runs[0].font.name = 'Consolas'

    # 2.3 Decisiones
    doc.add_heading("2.3. Decisiones de Diseño Clave", level=2)

    decisiones = [
        ("Separar settings por entorno",
         "base.py + development.py + production.py para no mezclar configuraciones"),
        ("Usar signals para auditoría",
         "LeadLog registra cambios automáticamente sin código explícito en cada vista"),
        ("Validación en 3 capas",
         "Campo → Validador → Modelo.clean() para máxima seguridad"),
        ("Honeypot + Rate limiting",
         "Doble protección anti-spam sin CAPTCHA molesto"),
        ("Magic bytes para archivos",
         "Validar contenido real, no solo extensión"),
    ]

    for titulo, descripcion in decisiones:
        p = doc.add_paragraph()
        p.add_run(f"• {titulo}: ").bold = True
        p.add_run(descripcion)

    doc.add_page_break()

def agregar_desarrollo(doc):
    """Agrega la sección de desarrollo por fases."""
    doc.add_heading("3. Desarrollo del Código - Fases", level=1)

    # Fase 1-3
    doc.add_heading("3.1. Fases 1-3: Estructura Base", level=2)

    doc.add_paragraph("FASE 1: Configuración inicial del proyecto", style='List Bullet')
    fase1 = [
        "Creación del proyecto Django",
        "Configuración de settings dividido (base/dev/prod)",
        "Configuración de .gitignore y estructura de carpetas",
        "Instalación de dependencias base",
    ]
    for item in fase1:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_paragraph("FASE 2: Modelos de datos", style='List Bullet')
    fase2 = [
        "Modelo Lead con todos los campos necesarios",
        "Modelo LeadImage para adjuntos",
        "Modelo Budget para presupuestos",
        "Modelo LeadLog para auditoría",
        "Modelo Service para catálogo",
        "Modelo UserProfile para roles",
    ]
    for item in fase2:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_paragraph("FASE 3: Panel de administración", style='List Bullet')
    fase3 = [
        "LeadAdmin con fieldsets colapsables",
        "Inlines para imágenes, presupuestos, historial",
        "Badges de estado con colores",
        "Filtros y búsqueda avanzada",
        "Optimización de queries (select_related, prefetch_related)",
    ]
    for item in fase3:
        p = doc.add_paragraph(f"   - {item}")

    # Fase 4-6
    doc.add_heading("3.2. Fases 4-6: Frontend y CRM", level=2)

    doc.add_paragraph("FASE 4: Templates y páginas públicas", style='List Bullet')
    fase4 = [
        "Template base con Tailwind CSS",
        "Página de inicio (home)",
        "Página de servicios",
        "Página de proyectos",
        "Página sobre nosotros",
        "Páginas legales (privacidad, cookies, aviso legal)",
    ]
    for item in fase4:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_paragraph("FASE 5: Integración Vite", style='List Bullet')
    fase5 = [
        "Configuración de Vite como bundler",
        "Compilación de CSS con Tailwind",
        "Gestión de assets estáticos",
        "Hot reload en desarrollo",
    ]
    for item in fase5:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_paragraph("FASE 6: Reorganización de templates", style='List Bullet')
    fase6 = [
        "Estructura pages/, legal/, errors/, components/",
        "Componentes reutilizables",
        "Templates de email",
    ]
    for item in fase6:
        p = doc.add_paragraph(f"   - {item}")

    # Fase 7-9
    doc.add_heading("3.3. Fases 7-9: Formulario y Seguridad", level=2)

    doc.add_paragraph("FASE 7: Formulario de contacto", style='List Bullet')
    fase7 = [
        "LeadForm con validaciones",
        "Vista contact_us con flujo completo",
        "Conexión formulario → modelo Lead",
        "Subida de imágenes (máx 5)",
        "Checkbox RGPD obligatorio",
    ]
    for item in fase7:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_paragraph("FASE 8: Seguridad anti-spam", style='List Bullet')
    fase8 = [
        "Rate limiting con django-ratelimit (5/hora por IP)",
        "Honeypot field (campo oculto website_url)",
        "Validación de magic bytes en archivos",
        "Límites de tamaño (5MB imágenes, 10MB PDFs)",
    ]
    for item in fase8:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_paragraph("FASE 9: Notificaciones por email", style='List Bullet')
    fase9 = [
        "Notificación al admin cuando llega un lead",
        "Confirmación automática al cliente",
        "Templates HTML para emails",
        "Configuración SMTP para producción",
    ]
    for item in fase9:
        p = doc.add_paragraph(f"   - {item}")

    doc.add_page_break()

def agregar_filosofias(doc):
    """Agrega la sección de filosofías aplicadas."""
    doc.add_heading("4. Filosofías y Convenciones Aplicadas", level=1)

    # 4.1 Clean Code y SOLID
    doc.add_heading("4.1. Clean Code y SOLID", level=2)

    doc.add_paragraph("Principios Clean Code aplicados:")
    clean_code = [
        ("Nombres descriptivos", "get_client_ip(), check_honeypot(), notify_new_lead()"),
        ("Funciones pequeñas", "Cada función hace una sola cosa"),
        ("Comentarios útiles", "Explican el 'por qué', no el 'qué'"),
        ("Código autodocumentado", "El código se entiende sin comentarios"),
        ("Sin código muerto", "Eliminación de código comentado"),
    ]
    for titulo, ejemplo in clean_code:
        p = doc.add_paragraph()
        p.add_run(f"• {titulo}: ").bold = True
        p.add_run(ejemplo)

    doc.add_paragraph()
    doc.add_paragraph("Principios SOLID aplicados:")
    solid = [
        ("Single Responsibility", "Cada módulo tiene una responsabilidad (models, forms, validators)"),
        ("Open/Closed", "FORM_SECURITY extensible sin modificar código base"),
        ("Liskov Substitution", "UserAdmin extiende BaseUserAdmin correctamente"),
        ("Interface Segregation", "Forms separados de Models"),
        ("Dependency Inversion", "Signals desacoplan la lógica de auditoría"),
    ]
    for titulo, ejemplo in solid:
        p = doc.add_paragraph()
        p.add_run(f"• {titulo}: ").bold = True
        p.add_run(ejemplo)

    # 4.2 Otras filosofías
    doc.add_heading("4.2. Otras Filosofías Web Importantes", level=2)

    doc.add_paragraph(
        "Durante el análisis del proyecto, se identificaron 20 filosofías/convenciones "
        "adicionales importantes para desarrollo web profesional:"
    )

    filosofias = [
        ("DRY (Don't Repeat Yourself)", "✅ Implementado", "Configuración centralizada"),
        ("KISS (Keep It Simple)", "✅ Implementado", "Arquitectura monolítica simple"),
        ("YAGNI (You Aren't Gonna Need It)", "✅ Implementado", "Sin features innecesarios"),
        ("Twelve-Factor App", "⚠️ Parcial", "Falta completar CI/CD y logs"),
        ("Defense in Depth", "⚠️ Parcial", "Falta CSP headers"),
        ("Fail-Safe Design", "✅ Implementado", "DEBUG=False en producción"),
        ("Graceful Degradation", "✅ Implementado", "Emails fallan sin romper leads"),
        ("Separation of Concerns", "✅ Implementado", "Models/Forms/Views separados"),
        ("Convention over Configuration", "⚠️ Parcial", "Falta linting automático"),
        ("Progressive Enhancement", "⚠️ Parcial", "Falta validación JS"),
        ("Mobile First", "✅ Implementado", "Tailwind CSS responsive"),
        ("Accessibility (a11y)", "⚠️ Parcial", "Falta ARIA labels"),
        ("SEO Best Practices", "⚠️ Parcial", "Falta meta tags dinámicos"),
        ("Performance Budget", "❌ Falta", "Sin medición Lighthouse"),
        ("Infrastructure as Code", "❌ Falta", "Sin Docker"),
        ("GitOps / CI/CD", "❌ Falta", "Sin GitHub Actions"),
        ("Observability", "❌ Falta", "Sin Sentry"),
        ("Data Privacy by Design", "✅ Implementado", "RGPD checkbox y políticas"),
    ]

    # Crear tabla
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filosofía'
    hdr_cells[1].text = 'Estado'
    hdr_cells[2].text = 'Notas'

    for nombre, estado, notas in filosofias:
        row_cells = table.add_row().cells
        row_cells[0].text = nombre
        row_cells[1].text = estado
        row_cells[2].text = notas

    doc.add_page_break()

def agregar_documentacion(doc):
    """Agrega la sección de documentación."""
    doc.add_heading("5. Documentación y Organización", level=1)

    # 5.1 Estructura
    doc.add_heading("5.1. Estructura de Documentación", level=2)

    doc.add_paragraph("Documentos creados durante el proyecto:")

    docs_list = [
        ("docs/DEPLOY_GUIDE.md", "Guía paso a paso para desplegar en producción"),
        ("docs/INFRAESTRUCTURA.md", "Análisis de opciones de hosting y costes"),
        ("docs/ANALISIS_FILOSOFIAS_WEB.md", "Análisis de filosofías aplicables"),
        ("SEED_DATABASE.md", "Instrucciones para poblar datos de prueba"),
        (".env.example", "Plantilla de variables de entorno"),
        ("requirements/*.txt", "Dependencias separadas por entorno"),
    ]

    for archivo, descripcion in docs_list:
        p = doc.add_paragraph()
        run = p.add_run(archivo)
        run.font.name = 'Consolas'
        run.bold = True
        p.add_run(f" - {descripcion}")

    # 5.2 Comentarios
    doc.add_heading("5.2. Comentarios en Código", level=2)

    doc.add_paragraph(
        "Se aplicó un estándar de documentación consistente en todos los archivos Python:"
    )

    doc.add_paragraph("Estructura de header en cada archivo:")
    code = '''"""
===============================================================================
ARCHIVO: ruta/al/archivo.py
PROYECTO: Arynstal - Sistema CRM para gestión de instalaciones y reformas
AUTOR: @cgvrzon
===============================================================================

DESCRIPCIÓN:
    Qué hace este archivo

FUNCIONES PRINCIPALES:
    - Lista de funciones importantes

FLUJO EN LA APLICACIÓN:
    1. Cómo se integra con el resto
===============================================================================
"""'''

    p = doc.add_paragraph(code)
    p.runs[0].font.name = 'Consolas'
    p.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("Estructura de docstring para funciones:")
    code2 = '''def contact_us(request):
    """
    Página de contacto con formulario para crear Leads.

    PROPÓSITO:
        Vista central del CRM. Procesa solicitudes de clientes.

    PARÁMETROS:
        request: Objeto HttpRequest de Django.

    RETORNA:
        HttpResponse: Página renderizada o redirect.

    FLUJO:
        1. Verificar rate limiting
        2. Verificar honeypot
        3. Validar formulario
        4. Crear Lead
        5. Enviar notificaciones
        6. Redirect con mensaje de éxito
    """'''

    p = doc.add_paragraph(code2)
    p.runs[0].font.name = 'Consolas'
    p.runs[0].font.size = Pt(9)

    doc.add_page_break()

def agregar_plan_produccion(doc):
    """Agrega el plan de producción."""
    doc.add_heading("6. Plan de Producción - Ruta Segura", level=1)

    doc.add_paragraph(
        "Se eligió la 'Ruta Segura' para el despliegue, que equilibra velocidad con "
        "profesionalismo. Esta ruta toma aproximadamente 1 semana e incluye las bases "
        "necesarias para un proyecto profesional."
    )

    # 6.1 Tareas
    doc.add_heading("6.1. Tareas Críticas a Implementar", level=2)

    tareas = [
        ("CI/CD con GitHub Actions", "1 día", "CRÍTICO",
         "Tests automáticos en cada push para prevenir bugs"),
        ("Integración Sentry", "2 horas", "CRÍTICO",
         "Monitoreo de errores en tiempo real"),
        ("CSP Headers", "2 horas", "CRÍTICO",
         "Protección contra XSS y inyección de scripts"),
        ("README.md profesional", "1 hora", "CRÍTICO",
         "Documentación de entrada al proyecto"),
        ("Meta tags SEO", "3 horas", "IMPORTANTE",
         "Visibilidad en Google desde el día 1"),
        ("Health check endpoint", "1 hora", "IMPORTANTE",
         "Monitoreo de disponibilidad"),
        ("Tests de integración", "2 días", "IMPORTANTE",
         "Cobertura de vistas completas"),
        ("2FA en admin", "3 horas", "IMPORTANTE",
         "Seguridad del panel de administración"),
    ]

    # Tabla de tareas
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tarea'
    hdr_cells[1].text = 'Tiempo'
    hdr_cells[2].text = 'Prioridad'
    hdr_cells[3].text = 'Justificación'

    for tarea, tiempo, prioridad, justificacion in tareas:
        row_cells = table.add_row().cells
        row_cells[0].text = tarea
        row_cells[1].text = tiempo
        row_cells[2].text = prioridad
        row_cells[3].text = justificacion

    # 6.2 Cronograma
    doc.add_heading("6.2. Cronograma Propuesto", level=2)

    cronograma = [
        ("Día 1", "CI/CD", "Configurar GitHub Actions con tests y linting"),
        ("Día 2", "Seguridad", "CSP headers + Sentry integration"),
        ("Día 3", "Documentación", "README.md + health check endpoint"),
        ("Día 4", "SEO", "Meta tags dinámicos + Open Graph"),
        ("Día 5", "Testing", "Tests de integración para vistas principales"),
        ("Día 6", "Admin", "2FA + revisión de seguridad"),
        ("Día 7", "Deploy", "Despliegue final + verificación"),
    ]

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Día'
    hdr_cells[1].text = 'Foco'
    hdr_cells[2].text = 'Tareas'

    for dia, foco, tareas_dia in cronograma:
        row_cells = table2.add_row().cells
        row_cells[0].text = dia
        row_cells[1].text = foco
        row_cells[2].text = tareas_dia

    doc.add_page_break()

def agregar_lecciones(doc):
    """Agrega las lecciones aprendidas."""
    doc.add_heading("7. Lecciones Aprendidas", level=1)

    lecciones = [
        ("Documentar desde el inicio",
         "La documentación durante el desarrollo es mucho más fácil que documentar "
         "después. Los comentarios en código ayudan a recordar decisiones."),

        ("Separar configuración por entorno",
         "Tener base.py + development.py + production.py evita muchos problemas "
         "y errores de configuración."),

        ("Validación en múltiples capas",
         "No confiar solo en la validación del frontend o del formulario. "
         "Validar en modelo también protege contra bugs y ataques."),

        ("Signals para efectos secundarios",
         "Usar signals para auditoría mantiene el código limpio y desacoplado. "
         "No hay que recordar añadir logs en cada vista."),

        ("Seguridad como diseño, no parche",
         "Integrar seguridad desde el inicio (CSRF, rate limiting, honeypot) "
         "es más fácil que añadirla después."),

        ("Testing es inversión, no gasto",
         "Los tests unitarios detectan problemas antes de que lleguen a producción. "
         "El tiempo invertido se recupera en debugging ahorrado."),

        ("El admin de Django es poderoso",
         "Personalizar el admin ahorra mucho tiempo vs crear interfaces custom. "
         "Fieldsets, inlines, badges hacen un admin profesional."),

        ("Git commits descriptivos",
         "Commits como 'FASE 7: Formulario de contacto conectado al modelo Lead' "
         "ayudan a entender la evolución del proyecto."),
    ]

    for titulo, descripcion in lecciones:
        p = doc.add_paragraph()
        p.add_run(f"• {titulo}").bold = True
        doc.add_paragraph(descripcion)
        doc.add_paragraph()

    doc.add_page_break()

def agregar_proximos_pasos(doc):
    """Agrega los próximos pasos."""
    doc.add_heading("8. Próximos Pasos", level=1)

    doc.add_paragraph(
        "Una vez completada la Ruta Segura, estos son los pasos futuros recomendados:"
    )

    doc.add_heading("Corto plazo (1-2 semanas post-lanzamiento):", level=2)
    corto = [
        "Monitorear errores en Sentry y corregir issues",
        "Analizar tráfico con Google Analytics",
        "Optimizar según métricas de Lighthouse",
        "Recoger feedback de usuarios reales",
    ]
    for item in corto:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Medio plazo (1-2 meses):", level=2)
    medio = [
        "Dockerizar la aplicación para reproducibilidad",
        "Implementar cache con Redis",
        "Añadir CDN para assets estáticos",
        "Crear dashboard de métricas de negocio",
        "Implementar exportación de leads a CSV/Excel",
    ]
    for item in medio:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Largo plazo (3-6 meses):", level=2)
    largo = [
        "API REST para posible app móvil",
        "Integración con WhatsApp Business API",
        "Sistema de plantillas de presupuesto",
        "Automatizaciones (emails de seguimiento)",
        "Multi-idioma si se expande a otros mercados",
    ]
    for item in largo:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # Cierre
    cierre = doc.add_paragraph()
    cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cierre.add_run("─" * 40)

    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.add_run("\nDocumento generado automáticamente\n").italic = True
    final.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}\n").italic = True
    final.add_run("Proyecto Arynstal - @cgvrzon").italic = True

def main():
    """Función principal que genera el documento."""
    print("Generando documento Word...")

    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Agregar contenido
    crear_estilos(doc)
    agregar_portada(doc)
    agregar_indice(doc)
    agregar_introduccion(doc)
    agregar_analisis(doc)
    agregar_desarrollo(doc)
    agregar_filosofias(doc)
    agregar_documentacion(doc)
    agregar_plan_produccion(doc)
    agregar_lecciones(doc)
    agregar_proximos_pasos(doc)

    # Guardar
    output_path = '/home/carlos/documentos/ARYNSTAL_Documentacion_Proceso_Desarrollo.docx'
    doc.save(output_path)
    print(f"✅ Documento guardado en: {output_path}")

    return output_path

if __name__ == '__main__':
    main()
